"""Tests for smart_filler.py — extract_docx_structure, format_structure_for_llm,
validate_fill_plan, execute_fill_plan, and _build_column_header_map."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from resume_tailor.templates.smart_filler import (
    _build_column_header_map,
    execute_fill_plan,
    extract_docx_structure,
    format_structure_for_llm,
    validate_fill_plan,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _create_docx_paragraphs(path: Path, texts: list[str]) -> Path:
    """Create a DOCX with only paragraphs."""
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


def _create_docx_with_table(
    path: Path,
    rows: int = 3,
    cols: int = 3,
    headers: list[str] | None = None,
) -> Path:
    """Create a DOCX with one paragraph and one table."""
    doc = Document()
    doc.add_paragraph("Test Document")
    table = doc.add_table(rows=rows, cols=cols)
    if headers:
        for i, h in enumerate(headers):
            if i < cols:
                table.rows[0].cells[i].text = h
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# extract_docx_structure tests
# ---------------------------------------------------------------------------

class TestExtractDocxStructure:
    def test_extract_paragraphs_count(self, tmp_path):
        """Structure contains one entry per non-empty paragraph."""
        path = tmp_path / "doc.docx"
        _create_docx_paragraphs(path, ["첫 번째 문단", "두 번째 문단"])

        structure = extract_docx_structure(path)

        assert len(structure["paragraphs"]) == 2

    def test_extract_paragraphs_text(self, tmp_path):
        """Paragraph entries contain the original text."""
        path = tmp_path / "doc.docx"
        _create_docx_paragraphs(path, ["안녕하세요"])

        structure = extract_docx_structure(path)

        assert structure["paragraphs"][0]["text"] == "안녕하세요"

    def test_extract_table_count(self, tmp_path):
        """Structure contains exactly one table when the DOCX has one table."""
        path = tmp_path / "doc.docx"
        _create_docx_with_table(path, rows=3, cols=3)

        structure = extract_docx_structure(path)

        assert len(structure["tables"]) == 1

    def test_extract_table_dimensions(self, tmp_path):
        """Table entry reports the correct row and column counts."""
        path = tmp_path / "doc.docx"
        _create_docx_with_table(path, rows=4, cols=2)

        structure = extract_docx_structure(path)
        table_info = structure["tables"][0]

        assert table_info["rows"] == 4
        assert table_info["cols"] == 2

    def test_extract_structure_keys(self, tmp_path):
        """Top-level structure always has 'paragraphs' and 'tables' keys."""
        path = tmp_path / "doc.docx"
        _create_docx_paragraphs(path, ["text"])

        structure = extract_docx_structure(path)

        assert "paragraphs" in structure
        assert "tables" in structure


# ---------------------------------------------------------------------------
# format_structure_for_llm tests
# ---------------------------------------------------------------------------

class TestFormatStructureForLlm:
    def test_returns_non_empty_string(self, tmp_path):
        """format_structure_for_llm returns a non-empty string."""
        path = tmp_path / "doc.docx"
        _create_docx_paragraphs(path, ["내용"])
        structure = extract_docx_structure(path)

        result = format_structure_for_llm(structure)

        assert isinstance(result, str)
        assert len(result) > 0

    def test_contains_paragraph_marker(self, tmp_path):
        """Output contains '문단' section header for documents with paragraphs."""
        path = tmp_path / "doc.docx"
        _create_docx_paragraphs(path, ["내용"])
        structure = extract_docx_structure(path)

        result = format_structure_for_llm(structure)

        assert "문단" in result

    def test_contains_table_marker(self, tmp_path):
        """Output contains '표' section header for documents with tables."""
        path = tmp_path / "doc.docx"
        _create_docx_with_table(path, rows=2, cols=2)
        structure = extract_docx_structure(path)

        result = format_structure_for_llm(structure)

        assert "표" in result

    def test_empty_structure_returns_empty_string(self):
        """An empty structure (no paragraphs, no tables) returns an empty string."""
        structure = {"paragraphs": [], "tables": []}
        result = format_structure_for_llm(structure)
        assert result == ""


# ---------------------------------------------------------------------------
# validate_fill_plan tests
# ---------------------------------------------------------------------------

class TestValidateFillPlan:
    def _make_structure_with_table(
        self,
        rows: int = 5,
        cols: int = 3,
        header_rows: list[int] | None = None,
    ) -> dict:
        """Build a minimal structure dict for a single table."""
        header_rows = header_rows or [0]

        header_cells = [{"col": c, "text": f"헤더{c}"} for c in range(cols)]
        data_rows = []
        for r in range(1, rows):
            data_rows.append({
                "row": r,
                "cells": [{"col": c, "empty": True} for c in range(cols)],
            })

        return {
            "paragraphs": [],
            "tables": [
                {
                    "idx": 0,
                    "rows": rows,
                    "cols": cols,
                    "header_rows": [{"row": r, "cells": header_cells} for r in header_rows],
                    "data_rows": data_rows,
                }
            ],
        }

    def test_valid_plan_returns_no_errors(self):
        """A well-formed plan targeting an existing data row returns no errors."""
        structure = self._make_structure_with_table(rows=5, cols=3)
        plan = {
            "fill_plan": [
                {
                    "target": "table",
                    "table_idx": 0,
                    "row": 2,
                    "fills": [{"col": 0, "value": "내용"}],
                }
            ]
        }

        errors = validate_fill_plan(structure, plan)

        assert errors == []

    def test_invalid_table_idx_returns_error(self):
        """Targeting a table_idx that does not exist produces an error."""
        structure = self._make_structure_with_table(rows=3, cols=2)
        plan = {
            "fill_plan": [
                {
                    "target": "table",
                    "table_idx": 5,
                    "row": 1,
                    "fills": [{"col": 0, "value": "값"}],
                }
            ]
        }

        errors = validate_fill_plan(structure, plan)

        assert len(errors) >= 1
        assert any("table_idx" in e or "범위" in e for e in errors)

    def test_header_row_target_returns_error(self):
        """Attempting to fill a header row produces an error."""
        structure = self._make_structure_with_table(rows=5, cols=3, header_rows=[0])
        plan = {
            "fill_plan": [
                {
                    "target": "table",
                    "table_idx": 0,
                    "row": 0,  # row 0 is a header row
                    "fills": [{"col": 1, "value": "값"}],
                }
            ]
        }

        errors = validate_fill_plan(structure, plan)

        assert len(errors) >= 1
        assert any("헤더" in e for e in errors)

    def test_empty_plan_returns_no_errors(self):
        """An empty fill_plan list produces no errors."""
        structure = self._make_structure_with_table()
        plan = {"fill_plan": []}

        errors = validate_fill_plan(structure, plan)

        assert errors == []


# ---------------------------------------------------------------------------
# execute_fill_plan tests
# ---------------------------------------------------------------------------

class TestExecuteFillPlan:
    def test_execute_fill_plan_fills_table_cell(self, tmp_path):
        """execute_fill_plan writes the specified value into the correct table cell."""
        src = tmp_path / "template.docx"
        _create_docx_with_table(src, rows=3, cols=3, headers=["A", "B", "C"])

        out = tmp_path / "filled.docx"
        plan = {
            "fill_plan": [
                {
                    "target": "table",
                    "table_idx": 0,
                    "row": 1,
                    "fills": [{"col": 0, "value": "채워진값"}],
                }
            ]
        }

        result = execute_fill_plan(src, plan, out)

        assert out.exists()
        doc = Document(str(out))
        cell_text = doc.tables[0].rows[1].cells[0].text
        assert "채워진값" in cell_text

    def test_execute_fill_plan_returns_output_path(self, tmp_path):
        """execute_fill_plan returns a Path pointing to the output file."""
        src = tmp_path / "template.docx"
        _create_docx_with_table(src, rows=2, cols=2)

        out = tmp_path / "filled.docx"
        plan = {"fill_plan": []}

        result = execute_fill_plan(src, plan, out)

        assert isinstance(result, Path)
        assert result == out


# ---------------------------------------------------------------------------
# _build_column_header_map tests
# ---------------------------------------------------------------------------

class TestBuildColumnHeaderMap:
    def test_returns_correct_mapping(self):
        """_build_column_header_map maps column indices to header text."""
        table_info = {
            "header_rows": [
                {
                    "row": 0,
                    "cells": [
                        {"col": 0, "text": "기간"},
                        {"col": 1, "text": "회사"},
                        {"col": 2, "text": "직위"},
                    ],
                }
            ],
            "data_rows": [],
        }

        result = _build_column_header_map(table_info)

        assert result == {0: "기간", 1: "회사", 2: "직위"}

    def test_empty_header_rows_returns_empty_map(self):
        """Table without header rows produces an empty mapping."""
        table_info = {"header_rows": [], "data_rows": []}

        result = _build_column_header_map(table_info)

        assert result == {}

    def test_ignores_empty_cell_text(self):
        """Cells with empty text are not included in the mapping."""
        table_info = {
            "header_rows": [
                {
                    "row": 0,
                    "cells": [
                        {"col": 0, "text": "이름"},
                        {"col": 1, "text": ""},
                        {"col": 2, "text": "직급"},
                    ],
                }
            ],
            "data_rows": [],
        }

        result = _build_column_header_map(table_info)

        assert 0 in result
        assert 2 in result
        assert 1 not in result
