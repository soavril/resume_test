"""Detailed tests for docx_renderer.py — generate_docx, fill_docx_template,
list_docx_placeholders, and _md_to_plain."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from resume_tailor.templates.docx_renderer import (
    _md_to_plain,
    fill_docx_template,
    generate_docx,
    list_docx_placeholders,
)


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------

def _create_template(path: Path, texts: list[str]) -> Path:
    """Create a simple DOCX with the given paragraph texts."""
    doc = Document()
    for text in texts:
        doc.add_paragraph(text)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# generate_docx tests
# ---------------------------------------------------------------------------

class TestGenerateDocx:
    def test_generate_docx_creates_file(self, tmp_path, sample_tailored_resume):
        """generate_docx creates the output file on disk."""
        out = tmp_path / "out.docx"
        generate_docx(sample_tailored_resume, out)
        assert out.exists()

    def test_generate_docx_returns_path(self, tmp_path, sample_tailored_resume):
        """generate_docx returns a Path object pointing to the file."""
        out = tmp_path / "out.docx"
        result = generate_docx(sample_tailored_resume, out)
        assert isinstance(result, Path)
        assert result == out

    def test_generate_docx_contains_sections(self, tmp_path, sample_tailored_resume):
        """Generated DOCX paragraphs include every section label from the resume."""
        out = tmp_path / "out.docx"
        generate_docx(sample_tailored_resume, out)

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)

        for section in sample_tailored_resume.sections:
            assert section.label in all_text, (
                f"Section label '{section.label}' not found in generated DOCX"
            )

    def test_generate_docx_creates_parent_dirs(self, tmp_path, sample_tailored_resume):
        """generate_docx creates missing parent directories."""
        out = tmp_path / "nested" / "deep" / "out.docx"
        generate_docx(sample_tailored_resume, out)
        assert out.exists()


# ---------------------------------------------------------------------------
# list_docx_placeholders tests
# ---------------------------------------------------------------------------

class TestListDocxPlaceholders:
    def test_list_docx_placeholders_finds_markers(self, tmp_path):
        """Placeholders {{이름}} and {{경력}} are detected and returned sorted."""
        tpl = tmp_path / "template.docx"
        _create_template(tpl, ["안녕하세요 {{이름}}님", "경력: {{경력}}"])

        result = list_docx_placeholders(tpl)

        assert result == ["경력", "이름"]

    def test_list_docx_placeholders_empty_template(self, tmp_path):
        """Template with no placeholders returns an empty list."""
        tpl = tmp_path / "empty.docx"
        _create_template(tpl, ["No placeholders here.", "Plain text."])

        result = list_docx_placeholders(tpl)

        assert result == []

    def test_list_docx_placeholders_deduplicates(self, tmp_path):
        """The same placeholder appearing twice is returned only once."""
        tpl = tmp_path / "dup.docx"
        _create_template(tpl, ["{{이름}} 위", "{{이름}} 아래"])

        result = list_docx_placeholders(tpl)

        assert result.count("이름") == 1


# ---------------------------------------------------------------------------
# fill_docx_template tests
# ---------------------------------------------------------------------------

class TestFillDocxTemplate:
    def test_fill_docx_replaces_section_by_id(self, tmp_path, sample_tailored_resume):
        """{{summary}} placeholder is replaced with the summary section content."""
        tpl = tmp_path / "tpl.docx"
        _create_template(tpl, ["자기소개: {{summary}}"])

        out = tmp_path / "filled.docx"
        fill_docx_template(tpl, sample_tailored_resume, out)

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # The placeholder must have been removed
        assert "{{summary}}" not in all_text
        # Section content must appear (plain text; markdown stripped)
        assert "대규모 트래픽" in all_text

    def test_fill_docx_with_extra_vars(self, tmp_path, sample_tailored_resume):
        """extra_vars dict values replace matching {{placeholder}} markers."""
        tpl = tmp_path / "tpl.docx"
        _create_template(tpl, ["커스텀 필드: {{custom_field}}"])

        out = tmp_path / "filled.docx"
        fill_docx_template(
            tpl,
            sample_tailored_resume,
            out,
            extra_vars={"custom_field": "특별한 값"},
        )

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "{{custom_field}}" not in all_text
        assert "특별한 값" in all_text

    def test_fill_docx_returns_output_path(self, tmp_path, sample_tailored_resume):
        """fill_docx_template returns the output path as a Path object."""
        tpl = tmp_path / "tpl.docx"
        _create_template(tpl, ["{{summary}}"])

        out = tmp_path / "filled.docx"
        result = fill_docx_template(tpl, sample_tailored_resume, out)

        assert isinstance(result, Path)
        assert result == out


# ---------------------------------------------------------------------------
# _md_to_plain tests
# ---------------------------------------------------------------------------

class TestMdToPlain:
    def test_strips_headers(self):
        """Markdown headers (## Title) are converted to plain title text."""
        result = _md_to_plain("## Title\nText")
        assert result == "Title\nText"

    def test_strips_h1_header(self):
        """H1 header prefix is removed."""
        result = _md_to_plain("# 홍길동")
        assert result == "홍길동"

    def test_strips_bold(self):
        """Bold markers (**text**) are removed, leaving plain text."""
        result = _md_to_plain("**bold** text")
        assert result == "bold text"

    def test_strips_italic(self):
        """Italic markers (*text*) are removed, leaving plain text."""
        result = _md_to_plain("*italic* word")
        assert result == "italic word"

    def test_strips_links(self):
        """Markdown links [text](url) are reduced to the link text."""
        result = _md_to_plain("[홈페이지](https://example.com)")
        assert result == "홈페이지"

    def test_strips_horizontal_rule(self):
        """Horizontal rules (---) are removed from the output."""
        result = _md_to_plain("line one\n---\nline two")
        assert "---" not in result

    def test_plain_text_unchanged(self):
        """Text without any markdown syntax is returned as-is."""
        plain = "일반 텍스트 내용입니다."
        result = _md_to_plain(plain)
        assert result == plain
