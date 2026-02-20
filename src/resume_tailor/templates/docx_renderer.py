"""DOCX output renderer — template-based fill + from-scratch generation."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from resume_tailor.models.resume import TailoredResume
from resume_tailor.parsers.resume_parser import EMOJI_PATTERN


# ---------------------------------------------------------------------------
# Mode 1: Template-based placeholder replacement
# ---------------------------------------------------------------------------

def fill_docx_template(
    template_path: str | Path,
    resume: TailoredResume,
    output_path: str | Path,
    extra_vars: dict[str, str] | None = None,
) -> Path:
    """Fill a .docx template by replacing {{placeholder}} markers.

    Supported placeholders (case-insensitive):
      - {{전체}} or {{full}}         → full_markdown (plain text)
      - {{섹션ID}} e.g. {{summary}} → section content by id
      - {{섹션라벨}} e.g. {{자기소개}} → section content by label
      - Any key from extra_vars      → custom value

    The replacement preserves the paragraph's existing formatting (font, size,
    color, bold, etc.) from the first run that contains the placeholder.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    # Build replacement map
    replacements = _build_replacement_map(resume, extra_vars)

    # Replace in paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    # Replace in headers/footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, replacements)

    doc.save(str(output_path))
    return output_path


def _build_replacement_map(
    resume: TailoredResume,
    extra_vars: dict[str, str] | None = None,
) -> dict[str, str]:
    """Build a {placeholder_key: content} map."""
    m: dict[str, str] = {}

    # Full resume
    m["전체"] = resume.full_markdown
    m["full"] = resume.full_markdown

    # By section id and label
    for section in resume.sections:
        m[section.id] = section.content
        m[section.label] = section.content

    if extra_vars:
        m.update(extra_vars)

    return m


def _replace_in_paragraph(para, replacements: dict[str, str]) -> None:
    """Replace {{key}} placeholders in a paragraph, preserving formatting.

    Handles the case where a placeholder may be split across multiple runs
    (e.g. Word sometimes splits {{자기소개}} into runs like "{{", "자기소개", "}}").
    """
    full_text = para.text
    if "{{" not in full_text:
        return

    # Find all placeholders in the combined text
    pattern = re.compile(r"\{\{(.+?)\}\}")
    matches = list(pattern.finditer(full_text))
    if not matches:
        return

    # Try simple run-level replacement first (placeholder in single run)
    for run in para.runs:
        for match in pattern.finditer(run.text):
            key = match.group(1).strip().lower()
            for rkey, rval in replacements.items():
                if rkey.lower() == key:
                    # Replace markdown content with plain text for docx
                    clean = _md_to_plain(rval)
                    run.text = run.text.replace(match.group(0), clean)
                    return

    # If placeholders span multiple runs, rebuild the paragraph
    for match in matches:
        key = match.group(1).strip().lower()
        for rkey, rval in replacements.items():
            if rkey.lower() == key:
                clean = _md_to_plain(rval)
                _rebuild_paragraph_with_replacement(para, match.group(0), clean)
                return


def _rebuild_paragraph_with_replacement(para, placeholder: str, replacement: str) -> None:
    """Rebuild paragraph runs when placeholder spans multiple runs."""
    full_text = "".join(run.text for run in para.runs)
    if placeholder not in full_text:
        return

    new_text = full_text.replace(placeholder, replacement)

    # Preserve formatting from the first run
    if para.runs:
        fmt_run = para.runs[0]
        # Clear all runs
        for run in para.runs:
            run.text = ""
        fmt_run.text = new_text


def _md_to_plain(md: str) -> str:
    """Convert simple markdown to plain text for DOCX embedding."""
    text = md
    # Remove markdown headers
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic markers
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    # Remove links [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Remove horizontal rules
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    # Remove emojis
    text = re.sub(EMOJI_PATTERN, "", text)
    # Remove excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def list_docx_placeholders(template_path: str | Path) -> list[str]:
    """Scan a .docx template and return all {{placeholder}} keys found."""
    doc = Document(str(template_path))
    placeholders = set()
    pattern = re.compile(r"\{\{(.+?)\}\}")

    for para in doc.paragraphs:
        for m in pattern.finditer(para.text):
            placeholders.add(m.group(1).strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for m in pattern.finditer(para.text):
                        placeholders.add(m.group(1).strip())

    return sorted(placeholders)


# ---------------------------------------------------------------------------
# Mode 2: Generate DOCX from scratch
# ---------------------------------------------------------------------------

def generate_docx(
    resume: TailoredResume,
    output_path: str | Path,
    title: str = "이력서",
) -> Path:
    """Generate a clean .docx from the resume sections (no template needed)."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    for section in resume.sections:
        _render_section(doc, section.label, section.content)

    doc.save(str(output_path))
    return output_path


def _render_section(doc: Document, label: str, content: str) -> None:
    """Render one resume section into the document."""
    # Section heading
    heading = doc.add_heading(label, level=2)
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Skip horizontal rules
        if re.match(r"^-{3,}\s*$", line):
            i += 1
            continue

        # Sub-heading (### )
        if line.startswith("### "):
            h = doc.add_heading(_strip_md_plain(line[4:]), level=3)
            h.runs[0].font.size = Pt(11)
            i += 1
            continue

        # Bullet point
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, line[2:])
            # Check for sub-bullets
            while i + 1 < len(lines) and lines[i + 1].startswith("  - "):
                i += 1
                sp = doc.add_paragraph(style="List Bullet 2")
                _add_rich_text(sp, lines[i].strip()[2:])
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_rich_text(p, line)
        i += 1


def _add_rich_text(paragraph, text: str) -> None:
    """Add text to a paragraph with bold/link formatting preserved."""
    # Remove emojis first
    text = _strip_emoji(text)
    # Remove link syntax [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

    # Split on bold markers and render with actual bold
    parts = re.split(r"(\*{2,3}.+?\*{2,3})", text)
    for part in parts:
        bold_match = re.match(r"\*{2,3}(.+?)\*{2,3}", part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)


def _strip_emoji(text: str) -> str:
    """Remove common emoji/icon characters."""
    return re.sub(EMOJI_PATTERN, "", text)


def _strip_md_plain(text: str) -> str:
    """Strip inline markdown formatting to plain text."""
    text = _strip_emoji(text)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text
