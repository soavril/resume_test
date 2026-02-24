"""LLM-based smart DOCX filler — handles any template structure.

Pipeline:
  1. Extract DOCX structure → compact JSON description
  2. LLM analyzes structure + resume → produces fill_plan
  3. Validate fill_plan against actual structure
  4. Retry with error feedback if validation fails
  5. Execute fill_plan on the DOCX
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from resume_tailor.clients.llm_client import LLMClient
from resume_tailor.models.resume import TailoredResume
from resume_tailor.utils.json_parser import extract_json

logger = logging.getLogger(__name__)

MAX_ADD_ROWS = 20

# ---------------------------------------------------------------------------
# Step 1: Extract DOCX structure
# ---------------------------------------------------------------------------


def extract_docx_structure(path: str | Path) -> dict:
    """Parse a DOCX file into a compact JSON structure description.

    Uses sequential unique-cell indices (0, 1, 2, ...) for consistency
    between extraction and execution. Each cell includes a column header
    mapping to help the LLM understand the table layout.
    """
    doc = Document(str(path))
    structure: dict = {"paragraphs": [], "tables": []}

    # Paragraphs (outside tables)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            structure["paragraphs"].append({
                "idx": i,
                "style": para.style.name,
                "text": text[:300],
            })

    # Tables
    for ti, table in enumerate(doc.tables):
        table_info: dict = {
            "idx": ti,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "header_rows": [],
            "data_rows": [],
        }

        # First pass: collect all rows to detect headers intelligently
        all_rows_data = []
        for ri, row in enumerate(table.rows):
            cells_info = []
            seen_ids = set()
            unique_idx = 0

            for i, cell in enumerate(row.cells):
                cid = id(cell._tc)
                if cid in seen_ids:
                    continue
                seen_ids.add(cid)

                # Count how many column slots this cell spans
                span = sum(1 for c in row.cells if id(c._tc) == cid)

                text = cell.text.strip()
                cell_data = {
                    "col": unique_idx,
                    "text": text[:300] if text else "",
                    "grid_start": i,
                }
                if span > 1:
                    cell_data["span"] = span
                if not text:
                    cell_data["empty"] = True

                cells_info.append(cell_data)
                unique_idx += 1

            all_rows_data.append({"row": ri, "cells": cells_info})

        # Second pass: classify rows as header or data
        for row_data in all_rows_data:
            ri = row_data["row"]
            cells = row_data["cells"]
            non_empty = [c for c in cells if not c.get("empty")]
            all_empty = len(non_empty) == 0

            if all_empty:
                table_info["data_rows"].append(row_data)
            elif _is_header_row(ri, cells, all_rows_data):
                table_info["header_rows"].append(row_data)
            else:
                table_info["data_rows"].append(row_data)

        structure["tables"].append(table_info)

    return structure


def _is_header_row(ri: int, cells: list[dict], all_rows: list[dict]) -> bool:
    """Detect header rows using content patterns rather than position alone.

    A row is a header if:
    - It's in the first 2 rows AND most cells are filled (not a data input row), OR
    - It has mostly non-empty cells and the next row has mostly empty cells
      (section header pattern), OR
    - Its cell texts look like field labels AND most cells are filled
    """
    non_empty = [c for c in cells if not c.get("empty")]
    non_empty_ratio = len(non_empty) / max(len(cells), 1)
    empty_ratio = 1 - non_empty_ratio

    # First 2 rows with content are almost always headers,
    # but rows with many empty cells may be data input rows (label + empty pattern)
    if ri < 2 and non_empty_ratio > 0.3 and empty_ratio < 0.3:
        return True

    # Check if this looks like a section header (most cells filled with labels)
    if non_empty_ratio >= 0.5:
        label_keywords = {"기간", "근무", "회사", "직위", "직급", "담당", "업무",
                          "학교", "학력", "전공", "기술", "자격", "수상", "어학",
                          "년", "월", "일", "시작", "종료", "성명", "생년월일",
                          "주소", "연락처", "이메일", "성별"}
        label_like = sum(
            1 for c in non_empty
            if any(kw in c["text"] for kw in label_keywords)
        )
        # Only treat as header if most cells are filled (not label + empty input pattern)
        if label_like >= 2 and empty_ratio < 0.3:
            return True

        # Check if next row is mostly empty (header → data pattern)
        if ri + 1 < len(all_rows):
            next_cells = all_rows[ri + 1]["cells"]
            next_empty_ratio = sum(
                1 for c in next_cells if c.get("empty")
            ) / max(len(next_cells), 1)
            if next_empty_ratio >= 0.5 and non_empty_ratio >= 0.6:
                return True

    return False


def _build_column_header_map(table_info: dict) -> dict[int, str]:
    """Build a mapping from column index to header text.

    Uses the first data row's grid_start values to map data-row TC indices
    to their corresponding headers.  This avoids mismatches when header
    rows and data rows have different TC counts due to cell merging.
    Falls back to header TC indices when no data rows exist.
    """
    data_rows = table_info.get("data_rows", [])
    header_rows = table_info.get("header_rows", [])

    if data_rows and header_rows:
        # Use grid_start-based mapping from the first data row
        return _map_row_to_headers(data_rows[0], header_rows)

    # Fallback: use header TC indices directly
    col_headers: dict[int, str] = {}
    for hr in header_rows:
        for cell in hr["cells"]:
            col = cell["col"]
            text = cell.get("text", "").strip()
            if text and col not in col_headers:
                col_headers[col] = text
    return col_headers


def _map_row_to_headers(
    data_row: dict, header_rows: list[dict],
) -> dict[int, str]:
    """Map data-row TC indices to header labels via proportional center alignment.

    When header and data rows have different cell merging (different TC counts),
    grid_start-based mapping shifts columns.  Instead, we compute each cell's
    proportional center (center position / row total grid width) and match
    data cells to the header cell with the closest proportional center.

    Only uses the nearest preceding header row to avoid cross-section mapping
    in multi-section tables (e.g. 학력/병역/경력 in one table).
    """
    data_row_num = data_row["row"]

    # Only consider header rows that PRECEDE this data row
    preceding = [hr for hr in header_rows if hr["row"] < data_row_num]
    if not preceding:
        return {}

    # Use only the nearest preceding header row
    nearest_row_num = max(hr["row"] for hr in preceding)
    nearest_headers = [hr for hr in preceding if hr["row"] == nearest_row_num]

    # Collect header cells that have text
    header_cells_with_text = [
        cell
        for hr in nearest_headers
        for cell in hr["cells"]
        if cell.get("text", "").strip()
    ]
    if not header_cells_with_text:
        return {}

    # Compute total grid width for header and data rows
    all_header_cells = [c for hr in nearest_headers for c in hr["cells"]]
    header_total = max(
        c.get("grid_start", c["col"]) + c.get("span", 1)
        for c in all_header_cells
    )
    data_total = max(
        c.get("grid_start", c["col"]) + c.get("span", 1)
        for c in data_row["cells"]
    )
    if header_total == 0 or data_total == 0:
        return {}

    # Build header proportional centers: [(center_ratio, text), ...]
    h_centers: list[tuple[float, str]] = []
    for cell in header_cells_with_text:
        gs = cell.get("grid_start", cell["col"])
        span = cell.get("span", 1)
        center = (gs + span / 2) / header_total
        h_centers.append((center, cell.get("text", "").strip()))

    # Map each data TC to the header with closest proportional center
    result: dict[int, str] = {}
    for cell in data_row["cells"]:
        tc_idx = cell["col"]
        gs = cell.get("grid_start", tc_idx)
        span = cell.get("span", 1)
        d_center = (gs + span / 2) / data_total

        best_dist = float("inf")
        best_text = None
        for h_center, h_text in h_centers:
            dist = abs(d_center - h_center)
            if dist < best_dist:
                best_dist = dist
                best_text = h_text

        if best_text:
            result[tc_idx] = best_text
    return result


def format_structure_for_llm(structure: dict) -> str:
    """Format the DOCX structure into a compact readable string for the LLM.

    Includes column-header mapping for each table to help the LLM
    understand what each column represents.
    """
    parts = []

    if structure["paragraphs"]:
        parts.append("=== 문단 (테이블 외부) ===")
        for p in structure["paragraphs"]:
            parts.append(f"  P[{p['idx']}] ({p['style']}): \"{p['text']}\"")

    for t in structure["tables"]:
        parts.append(f"\n=== 표 {t['idx']} ({t['rows']}행) ===")

        # Column-header mapping
        col_headers = _build_column_header_map(t)
        if col_headers:
            parts.append("  열-헤더 매핑 (col = 셀 순번):")
            for col_idx in sorted(col_headers.keys()):
                parts.append(f"    col{col_idx}: \"{col_headers[col_idx]}\"")

        if t["header_rows"]:
            parts.append("  헤더 행 (수정 금지):")
            for hr in t["header_rows"]:
                cells_str = " | ".join(
                    f"[col{c['col']}"
                    + (f",span{c['span']}" if c.get("span") else "")
                    + f"] \"{c['text']}\""
                    for c in hr["cells"]
                )
                parts.append(f"    Row {hr['row']}: {cells_str}")

        if t["data_rows"]:
            empty_rows = [r for r in t["data_rows"]
                          if all(c.get("empty") for c in r["cells"])]
            filled_rows = [r for r in t["data_rows"]
                           if not all(c.get("empty") for c in r["cells"])]

            if filled_rows:
                parts.append("  채워진 데이터 행:")
                for fr in filled_rows:
                    tc_count = len(fr["cells"])
                    filled_str = " | ".join(
                        f"[col{c['col']}] \"{c['text']}\""
                        for c in fr["cells"]
                        if not c.get("empty")
                    )
                    empty_cols = [
                        c["col"] for c in fr["cells"] if c.get("empty")
                    ]
                    empty_hint = ""
                    if empty_cols:
                        col_map = _map_row_to_headers(fr, t["header_rows"])
                        hints = [
                            f"col{c}(\"{col_map[c]}\")"
                            if c in col_map else f"col{c}"
                            for c in empty_cols
                        ]
                        empty_hint = f" ← 빈 셀: {', '.join(hints)}"
                    parts.append(
                        f"    Row {fr['row']} (col 범위: 0~{tc_count - 1}): "
                        f"{filled_str}{empty_hint}"
                    )

            if empty_rows:
                # Group by TC count since merged rows have different cell counts
                tc_groups: dict[int, list[int]] = {}
                for r in empty_rows:
                    tc = len(r["cells"])
                    tc_groups.setdefault(tc, []).append(r["row"])

                if len(tc_groups) == 1:
                    tc_count = next(iter(tc_groups))
                    first = empty_rows[0]["row"]
                    last = empty_rows[-1]["row"]
                    parts.append(
                        f"  빈 데이터 행: Row {first}~{last} "
                        f"({len(empty_rows)}행, col 범위: 0~{tc_count - 1})"
                    )
                    # Show header mapping for empty rows
                    col_map = _map_row_to_headers(
                        empty_rows[0], t["header_rows"],
                    )
                    if col_map:
                        mapping_parts = [
                            f"col{c}→\"{col_map[c]}\""
                            for c in sorted(col_map)
                        ]
                        parts.append(
                            f"    열 매핑: {', '.join(mapping_parts)}"
                        )
                else:
                    parts.append(
                        f"  빈 데이터 행 ({len(empty_rows)}행, "
                        f"행별 셀 수 다름 — 병합 구조):"
                    )
                    for tc_count in sorted(tc_groups):
                        rows = tc_groups[tc_count]
                        row_str = _format_row_ranges(rows)
                        parts.append(
                            f"    {row_str}: col 범위 0~{tc_count - 1}"
                        )
                        # Show header mapping for this row group
                        sample_row = next(
                            r for r in empty_rows
                            if len(r["cells"]) == tc_count
                        )
                        col_map = _map_row_to_headers(
                            sample_row, t["header_rows"],
                        )
                        if col_map:
                            mapping_parts = [
                                f"col{c}→\"{col_map[c]}\""
                                for c in sorted(col_map)
                            ]
                            parts.append(
                                f"      열 매핑: {', '.join(mapping_parts)}"
                            )

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Step 2: LLM produces fill_plan
# ---------------------------------------------------------------------------

def _format_row_ranges(rows: list[int]) -> str:
    """Format row numbers into compact ranges: [1,2,3,5,6] → 'Row 1~3, 5~6'."""
    if not rows:
        return ""
    ranges: list[str] = []
    start = end = rows[0]
    for r in rows[1:]:
        if r == end + 1:
            end = r
        else:
            ranges.append(f"{start}~{end}" if start != end else str(start))
            start = end = r
    ranges.append(f"{start}~{end}" if start != end else str(start))
    return "Row " + ", ".join(ranges)


ANALYZER_SYSTEM = """\
당신은 DOCX 이력서 양식 분석 전문가입니다.

주어진 DOCX 구조와 이력서 내용을 분석하여, 양식의 어느 위치에 어떤 내용을 채워야 하는지 계획을 세웁니다.

## 중요 규칙

1. **절대 이력서에 없는 정보를 만들어내지 마세요.** 회사명, 학교명, 수치, 날짜, 연락처 등은 반드시 이력서 원본에 있는 사실만 사용하세요. 정보가 부족하면 해당 칸을 비워두세요 (fill_plan에서 제외).
2. **모든 빈 셀을 최대한 채우세요.** 이력서에 해당 정보가 있으면 반드시 fill_plan에 포함하세요. "열 매핑"을 참고하여 각 col에 맞는 정보를 넣으세요. 채워진 데이터 행의 빈 셀("← 빈 셀" 표시)도 채울 수 있으면 채우세요. 학력, 경력, 병역, 외국어, 해외연수, 교내외활동, 자기소개서, 경력기술서, 희망연봉 등 모든 섹션을 빠짐없이 채우세요. 자리표시자('년 월' 등)가 있는 셀도 실제 데이터로 교체해야 합니다.
3. **각 셀에는 해당 열 헤더에 맞는 단일 정보만 넣으세요.** 여러 항목을 한 셀에 합치지 마세요. 예: "회사명" 열에는 회사명만, "직급" 열에는 직급만. 장소·기관·내용이 별도 열이면 각각 분리하세요.
4. **col 인덱스는 고유 셀 순번입니다 (0부터 시작).** 각 행마다 병합 패턴이 달라 셀 수가 다릅니다. 반드시 "col 범위: 0~N" 표시를 확인하고 그 범위 안에서만 col을 사용하세요.
5. **병합된 셀**: span이 표시된 셀은 여러 열을 차지하지만 하나의 col 인덱스(순번)로만 참조합니다. 예: col0이 span3이면 다음 셀은 col1입니다 (col3이 아닙니다!).
6. **헤더 행 수정 금지**: "헤더 행 (수정 금지)"으로 표시된 행의 Row 번호는 절대 fill_plan에 포함하지 마세요.
7. **빈 데이터 행과 채워진 행의 빈 셀** 모두 채울 수 있습니다.
8. 경력 항목은 **최근부터 역순**으로 배치합니다.
9. 날짜 열: 년/월/일이 개별 열이면 각각 분리해서 넣으세요.
10. 줄바꿈이 필요한 경우 \\n을 사용하세요 (실제 줄바꿈으로 변환됩니다).
11. 마크다운 서식(**, #, - 등)을 제거한 순수 텍스트로 작성하세요.
12. **자기소개서/경력기술서 셀**: 원본 질문이 있는 셀에는 질문을 유지하고 그 아래에 답변을 이어서 작성하세요 (질문\\n\\n답변 형식).
13. **자리표시자 vs 라벨+단위 서식 구분**:
    - '년    월', '년  월-  년  월' 같은 **날짜 자리표시자**는 실제 날짜로 **교체**하세요. 예: '년    월' → '2023 09', '년  월-  년  월' → '2019.01 - 2019.09'.
    - '연봉      만원' 같은 **라벨+단위 서식**이 셀에 이미 있으면 빈칸 부분만 채우세요. 예: '연봉      만원' → '연봉 7,500 만원'.
    - **빈 셀**인데 열 매핑 라벨이 '연봉 만원' 같은 서식이면, 값만 넣으세요 (서식 반복 금지). 예: 열 매핑 '연봉 만원' + 빈 셀 → '7,500' 또는 '협의 가능' (NOT '연봉 7,500 만원', NOT '연봉 협의 가능 만원').
    - **(한글)**, **(한자)**, **(영문)** 같은 라벨은 실제 데이터로 교체하세요. 예: '(한글)' → '최홍익'.
14. **중복 금지**: 같은 행의 여러 빈 셀에 동일한 값을 넣지 마세요. 각 빈 셀은 서로 다른 정보 항목입니다. 인접 라벨 텍스트(예: '성명', '(한글)', '주소')를 참고하여 각 셀에 맞는 값을 넣으세요.
15. **학교 유형 라벨 교체**: '고등학교', '전문대학', '대학교', '대학원(석사)', '대학원(박사)' 같은 학교 유형 텍스트가 학교명 열에 있으면, 이력서의 실제 학교명으로 **교체**하세요. 예: '대학교' → '서울과학기술대학교'. 이것은 채워진 셀이 아니라 자리표시자입니다.

## 예시

다음과 같은 한국 경력기술서 표 구조가 있을 때:

```
=== 표 0 (10행 x 8열) ===
  열-헤더 매핑:
    col0: "근무기간"
    col1: "근무기간"
    col2: "회사명"
    col3: "부서/직급"
    col4: "담당업무"
  헤더 행:
    Row 0: [col0,span2] "근무기간" | [col2] "회사명" | [col3] "부서/직급" | [col4] "담당업무"
    Row 1: [col0] "시작" | [col1] "종료" | [col2] "" | [col3] "" | [col4] ""
  빈 데이터 행: Row 2~9 (8행, 각 5열)
```

올바른 fill_plan:
```json
{
  "analysis": "경력기술서 표. 5열 구조: 시작일, 종료일, 회사명, 부서/직급, 담당업무.",
  "fill_plan": [
    {
      "target": "table",
      "table_idx": 0,
      "row": 2,
      "fills": [
        {"col": 0, "value": "2022.03"},
        {"col": 1, "value": "현재"},
        {"col": 2, "value": "삼성전자"},
        {"col": 3, "value": "개발팀/선임"},
        {"col": 4, "value": "백엔드 API 설계 및 개발\\n성능 최적화 30% 달성"}
      ]
    }
  ]
}
```

## 응답 형식

반드시 아래 JSON 형식으로 응답하세요:
{
  "analysis": "양식 구조에 대한 간단한 분석",
  "fill_plan": [
    {
      "target": "table",
      "table_idx": 0,
      "row": 2,
      "fills": [
        {"col": 0, "value": "2022"},
        {"col": 1, "value": "03"},
        {"col": 7, "value": "상세 내용..."}
      ]
    },
    {
      "target": "paragraph",
      "paragraph_idx": 3,
      "action": "replace",
      "value": "새 내용"
    },
    {
      "target": "paragraph",
      "after_paragraph_idx": 0,
      "action": "insert",
      "value": "추가할 내용"
    }
  ]
}"""


def _split_table_into_sections(table: dict) -> list[dict]:
    """Split a multi-header table into per-section mini-tables.

    Tables with ≤1 header row are returned as-is.
    Each section = one header row + its following data rows (until next header).
    """
    headers = table["header_rows"]
    if len(headers) <= 1:
        return [table]

    header_row_nums = sorted(hr["row"] for hr in headers)
    sections: list[dict] = []

    for i, h_row_num in enumerate(header_row_nums):
        next_h = (
            header_row_nums[i + 1]
            if i + 1 < len(header_row_nums)
            else float("inf")
        )
        section_headers = [hr for hr in headers if hr["row"] == h_row_num]
        section_data = [
            dr for dr in table["data_rows"]
            if h_row_num < dr["row"] < next_h
        ]

        if not section_data:
            continue

        section_rows = len(section_headers) + len(section_data)
        sections.append({
            "idx": table["idx"],
            "rows": section_rows,
            "header_rows": section_headers,
            "data_rows": section_data,
        })

    return sections


def _build_section_groups(structure: dict) -> list[dict]:
    """Split structure into groups for parallel LLM calls.

    Multi-header tables are split by section so each LLM call only sees
    one section's headers and data rows — preventing cross-section column
    confusion.  Simple tables are grouped together into a single call.
    """
    section_groups: list[dict] = []
    simple_tables: list[dict] = []

    for table in structure["tables"]:
        sections = _split_table_into_sections(table)
        if len(sections) <= 1:
            simple_tables.append(table)
        else:
            for section in sections:
                section_groups.append({
                    "tables": [section],
                    "paragraphs": [],
                })

    # All simple tables + paragraphs in one group
    if simple_tables or structure.get("paragraphs"):
        section_groups.insert(0, {
            "tables": simple_tables,
            "paragraphs": structure.get("paragraphs", []),
        })

    return section_groups


async def _analyze_section(
    llm: LLMClient,
    section_structure: dict,
    resume_content: str,
    model: str,
) -> dict:
    """Analyze a single section/group and produce a fill plan."""
    structure_text = format_structure_for_llm(section_structure)

    prompt = f"""다음 양식 구조를 분석하고, 이력서 내용을 채워넣을 계획을 세우세요.

## 양식 구조
{structure_text}

## 채울 이력서 내용
{resume_content}

양식의 열-헤더 매핑을 정확히 참고하여, 각 빈 칸에 어떤 내용을 넣을지 fill_plan JSON으로 응답하세요.
col 인덱스는 고유 셀 순번(0부터 시작)입니다."""

    return await llm.generate_json(
        prompt=prompt,
        system=ANALYZER_SYSTEM,
        model=model,
        max_tokens=8192,
    )


async def analyze_and_plan(
    llm: LLMClient,
    structure: dict,
    resume: TailoredResume,
    model: str = "claude-sonnet-4-5-20250929",
) -> dict:
    """LLM analyzes structure and produces a fill plan.

    Multi-header tables are split into sections and analyzed in parallel
    to prevent cross-section column confusion.
    """
    # Build resume content — fallback to sections if full_markdown is sparse
    resume_content = resume.full_markdown
    if not resume_content.strip() or len(resume_content.strip()) < 50:
        parts = [f"### {s.label}\n{s.content}" for s in resume.sections]
        resume_content = "\n\n".join(parts)

    groups = _build_section_groups(structure)

    if len(groups) <= 1:
        # Small template: single call (original path)
        return await _analyze_section(
            llm, groups[0] if groups else structure, resume_content, model,
        )

    # Parallel calls for each section group
    tasks = [
        _analyze_section(llm, group, resume_content, model)
        for group in groups
    ]
    results = await asyncio.gather(*tasks, return_exceptions=True)

    # Merge fill plans
    merged_plan: list[dict] = []
    analyses: list[str] = []
    for r in results:
        if isinstance(r, Exception):
            logger.warning("Section analysis failed: %s", r)
            continue
        analyses.append(r.get("analysis", ""))
        merged_plan.extend(r.get("fill_plan", []))

    return {
        "analysis": " | ".join(a for a in analyses if a),
        "fill_plan": merged_plan,
    }


# ---------------------------------------------------------------------------
# Step 2.5: Validate fill_plan
# ---------------------------------------------------------------------------


def validate_fill_plan(structure: dict, plan: dict) -> list[str]:
    """Validate a fill_plan against the actual DOCX structure.

    Returns a list of error messages. Empty list means valid.
    """
    errors = []
    fills = plan.get("fill_plan", [])
    tables = structure.get("tables", [])
    paragraphs = structure.get("paragraphs", [])

    # Track which cells are targeted (detect duplicates)
    seen_cells: set[tuple[int, int, int]] = set()  # (table_idx, row, col)

    # Collect header rows per table for checking
    header_rows_by_table: dict[int, set[int]] = {}
    for t in tables:
        header_rows_by_table[t["idx"]] = {hr["row"] for hr in t["header_rows"]}

    # Collect valid column indices per table row
    valid_cols_by_row: dict[tuple[int, int], set[int]] = {}
    for t in tables:
        for row_data in t["header_rows"] + t["data_rows"]:
            key = (t["idx"], row_data["row"])
            valid_cols_by_row[key] = {c["col"] for c in row_data["cells"]}

    for i, item in enumerate(fills):
        target = item.get("target")

        if target == "table":
            table_idx = item.get("table_idx", 0)
            row = item.get("row")

            # Check table_idx range
            if table_idx >= len(tables):
                errors.append(
                    f"fill_plan[{i}]: table_idx={table_idx} 범위 초과 "
                    f"(표 {len(tables)}개)"
                )
                continue

            if row is None:
                errors.append(f"fill_plan[{i}]: row가 지정되지 않음")
                continue

            # Check if row is a header row
            if row in header_rows_by_table.get(table_idx, set()):
                errors.append(
                    f"fill_plan[{i}]: Row {row}은 헤더 행이므로 수정 불가"
                )
                continue

            # Check row range (allow adding rows up to MAX_ADD_ROWS beyond existing)
            table = tables[table_idx]
            max_row = table["rows"] + MAX_ADD_ROWS
            if row >= max_row:
                errors.append(
                    f"fill_plan[{i}]: row={row} 범위 초과 "
                    f"(표 행 수: {table['rows']}, 최대 추가: {MAX_ADD_ROWS})"
                )
                continue

            # Validate column indices in fills
            valid_cols = valid_cols_by_row.get((table_idx, row))
            for fill in item.get("fills", []):
                col = fill.get("col")
                if col is None:
                    continue

                # Check duplicate cell targeting
                cell_key = (table_idx, row, col)
                if cell_key in seen_cells:
                    errors.append(
                        f"fill_plan[{i}]: 표{table_idx} Row{row} col{col} "
                        f"중복 채우기 감지"
                    )
                seen_cells.add(cell_key)

                # Check col validity if we know the row structure
                if valid_cols is not None and col not in valid_cols:
                    errors.append(
                        f"fill_plan[{i}]: 표{table_idx} Row{row}에 "
                        f"col{col}이 존재하지 않음 "
                        f"(유효: {sorted(valid_cols)})"
                    )

        elif target == "paragraph":
            action = item.get("action", "replace")
            if action == "replace":
                idx = item.get("paragraph_idx")
                if idx is not None and idx >= len(paragraphs):
                    errors.append(
                        f"fill_plan[{i}]: paragraph_idx={idx} 범위 초과"
                    )
            elif action == "insert":
                after_idx = item.get("after_paragraph_idx")
                if after_idx is not None and after_idx >= len(paragraphs):
                    errors.append(
                        f"fill_plan[{i}]: after_paragraph_idx={after_idx} "
                        f"범위 초과"
                    )

    return errors


# ---------------------------------------------------------------------------
# Step 3: Execute fill_plan
# ---------------------------------------------------------------------------


def execute_fill_plan(
    doc_path: str | Path, plan: dict, output_path: str | Path,
) -> Path:
    """Execute a fill plan on a DOCX document."""
    doc_path = Path(doc_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(doc_path))
    fills = plan.get("fill_plan", [])

    filled_count = 0
    failed_count = 0

    for item in fills:
        target = item.get("target")

        if target == "table":
            ok, fail = _execute_table_fill(doc, item)
            filled_count += ok
            failed_count += fail
        elif target == "paragraph":
            _execute_paragraph_fill(doc, item)
            filled_count += 1

    logger.info("DOCX fill complete: %d cells filled, %d failed",
                filled_count, failed_count)

    doc.save(str(output_path))
    return output_path


def _execute_table_fill(doc: Document, item: dict) -> tuple[int, int]:
    """Fill table cells according to the plan.

    Returns (filled_count, failed_count).
    """
    table_idx = item.get("table_idx", 0)
    row_idx = item.get("row")
    filled = 0
    failed = 0

    if table_idx >= len(doc.tables) or row_idx is None:
        return 0, 1

    table = doc.tables[table_idx]

    # Add rows if needed (with upper bound)
    rows_added = 0
    while row_idx >= len(table.rows) and rows_added < MAX_ADD_ROWS:
        _add_table_row(table)
        rows_added += 1

    if row_idx >= len(table.rows):
        logger.warning(
            "Row %d exceeds table size and MAX_ADD_ROWS (%d)", row_idx, MAX_ADD_ROWS
        )
        return 0, len(item.get("fills", []))

    row = table.rows[row_idx]

    # Build a map from unique-cell-index to actual cell objects
    # Uses SAME sequential indexing as extract_docx_structure
    unique_cells = _get_unique_cells_with_col(row)

    for fill in item.get("fills", []):
        col = fill.get("col")
        value = str(fill.get("value", ""))
        if col is None or not value:
            continue

        cell = unique_cells.get(col)
        if cell:
            _set_cell_text(cell, _strip_md(value))
            filled += 1
        else:
            logger.warning(
                "Table %d Row %d: col %d not found (valid: %s)",
                table_idx, row_idx, col, sorted(unique_cells.keys()),
            )
            failed += 1

    return filled, failed


def _execute_paragraph_fill(doc: Document, item: dict) -> None:
    """Replace or insert a paragraph."""
    action = item.get("action", "replace")
    value = _strip_md(str(item.get("value", "")))

    if action == "replace":
        idx = item.get("paragraph_idx")
        if idx is not None and idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            # Clear and set
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = value
            else:
                para.add_run(value)

    elif action == "insert":
        after_idx = item.get("after_paragraph_idx")
        if after_idx is not None and after_idx < len(doc.paragraphs):
            para = doc.paragraphs[after_idx]
            _insert_paragraph_after(para, value)


def _insert_paragraph_after(paragraph, text: str):
    """Insert a new paragraph after the given paragraph."""
    from copy import deepcopy

    new_p = deepcopy(paragraph._p)
    # Clear content
    for child in list(new_p):
        if child.tag.endswith("}r"):
            new_p.remove(child)

    paragraph._p.addnext(new_p)

    # Add text run
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    new_p.append(r)

    return new_p


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _get_unique_cells_with_col(row) -> dict:
    """Get a {col_index: cell} mapping of unique cells in a row.

    Uses sequential unique-cell indices (0, 1, 2, ...) matching
    the indexing in extract_docx_structure.
    """
    result = {}
    seen = set()
    unique_idx = 0
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result[unique_idx] = cell
            unique_idx += 1
    return result


def _set_cell_text(cell, text: str, font_size: int = 10) -> None:
    """Set cell text, preserving existing formatting.

    Handles multiline text by creating proper paragraph breaks.
    """
    # Capture existing format from first run
    existing_font_size = None
    existing_bold = None
    existing_font_name = None
    if cell.paragraphs and cell.paragraphs[0].runs:
        first_run = cell.paragraphs[0].runs[0]
        if first_run.font.size:
            existing_font_size = first_run.font.size
        existing_bold = first_run.font.bold
        existing_font_name = first_run.font.name

    # Clear all existing content
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""

    # Split on \n for multiline support
    lines = text.split("\n")

    if cell.paragraphs:
        # Use first paragraph for first line
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = lines[0]
        else:
            run = para.add_run(lines[0])
            _apply_font(run, existing_font_size, existing_bold,
                        existing_font_name, font_size)

        # Add additional paragraphs for remaining lines
        for line in lines[1:]:
            new_para = cell.add_paragraph()
            run = new_para.add_run(line)
            _apply_font(run, existing_font_size, existing_bold,
                        existing_font_name, font_size)
    else:
        # No existing paragraphs — create one
        para = cell.add_paragraph()
        run = para.add_run(lines[0])
        _apply_font(run, existing_font_size, existing_bold,
                    existing_font_name, font_size)
        for line in lines[1:]:
            new_para = cell.add_paragraph()
            run = new_para.add_run(line)
            _apply_font(run, existing_font_size, existing_bold,
                        existing_font_name, font_size)


def _apply_font(
    run, existing_size, existing_bold, existing_name, default_size: int,
) -> None:
    """Apply font formatting to a run, preferring existing style."""
    if existing_size:
        run.font.size = existing_size
    elif default_size:
        run.font.size = Pt(default_size)
    if existing_bold is not None:
        run.font.bold = existing_bold
    if existing_name:
        run.font.name = existing_name


def _add_table_row(table) -> None:
    """Add a new row copying format from the last row."""
    from copy import deepcopy

    last_row = table.rows[-1]
    new_tr = deepcopy(last_row._tr)
    for tc in new_tr.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    ):
        tc.text = ""
    table._tbl.append(new_tr)


def _strip_md(text: str) -> str:
    """Strip markdown formatting."""
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Step 4: Retry with validation errors
# ---------------------------------------------------------------------------


async def _retry_with_errors(
    llm: LLMClient,
    structure: dict,
    resume: TailoredResume,
    plan: dict,
    errors: list[str],
    model: str = "claude-sonnet-4-5-20250929",
) -> dict:
    """Ask LLM to fix the fill_plan based on validation errors."""
    structure_text = format_structure_for_llm(structure)
    errors_text = "\n".join(f"- {e}" for e in errors)

    # Build resume content — fallback to sections if full_markdown is sparse
    resume_content = resume.full_markdown
    if not resume_content.strip() or len(resume_content.strip()) < 50:
        parts = [f"### {s.label}\n{s.content}" for s in resume.sections]
        resume_content = "\n\n".join(parts)

    prompt = f"""이전 fill_plan에 다음 오류가 발견되었습니다:

{errors_text}

## DOCX 양식 구조
{structure_text}

## 이전 fill_plan
{json.dumps(plan, ensure_ascii=False, indent=2)}

## 이력서 내용
{resume_content}

오류를 수정한 새로운 fill_plan JSON으로 응답하세요.
col 인덱스는 고유 셀 순번(0부터 시작)이며, 열-헤더 매핑을 참고하세요."""

    return await llm.generate_json(
        prompt=prompt,
        system=ANALYZER_SYSTEM,
        model=model,
        max_tokens=16384,
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


async def smart_fill_docx(
    template_path: str | Path,
    resume: TailoredResume,
    output_path: str | Path,
    llm: LLMClient,
    max_attempts: int = 2,
    model: str = "claude-sonnet-4-5-20250929",
) -> Path:
    """Analyze any DOCX template and intelligently fill it with resume data.

    1. Extract DOCX structure
    2. LLM analyzes structure + resume → fill_plan
    3. Validate fill_plan
    4. Retry with error feedback if validation fails (up to max_attempts)
    5. Execute fill_plan on the DOCX
    """
    structure = extract_docx_structure(template_path)

    plan = await analyze_and_plan(llm, structure, resume, model=model)

    for attempt in range(max_attempts):
        errors = validate_fill_plan(structure, plan)
        if not errors:
            break
        logger.warning(
            "Fill plan validation failed (attempt %d/%d): %s",
            attempt + 1, max_attempts, errors,
        )
        if attempt < max_attempts - 1:
            plan = await _retry_with_errors(
                llm, structure, resume, plan, errors, model=model,
            )

    if errors:
        logger.warning("Proceeding with %d validation errors after "
                       "%d attempts", len(errors), max_attempts)

    return execute_fill_plan(template_path, plan, output_path)


def smart_fill_docx_sync(
    template_path: str | Path,
    resume: TailoredResume,
    output_path: str | Path,
    llm: LLMClient,
) -> Path:
    """Synchronous wrapper for smart_fill_docx."""
    return asyncio.run(
        smart_fill_docx(template_path, resume, output_path, llm)
    )
