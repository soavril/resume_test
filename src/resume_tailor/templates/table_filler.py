"""Fill table-based DOCX templates (경력기술서 등) with resume data.

Handles the common Korean 경력기술서 format:
  Row 0: Header "근무/프로젝트참여기간" (merged) | "주요 프로젝트"
  Row 1: 년 | 월 | 일 | ~ | 년 | 월 | 일 | (merged with header)
  Row 2+: Empty data rows to fill
"""

from __future__ import annotations

import asyncio
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt

from resume_tailor.clients.llm_client import LLMClient
from resume_tailor.models.resume import TailoredResume
from resume_tailor.utils.json_parser import extract_json

EXTRACT_PROMPT_SYSTEM = """\
당신은 이력서에서 경력 항목을 추출하는 전문가입니다.
이력서 텍스트에서 각 경력/프로젝트 항목을 아래 JSON 배열로 추출하세요.

각 항목:
{
  "start_year": "2022",
  "start_month": "03",
  "start_day": "01",
  "end_year": "2025",
  "end_month": "02",
  "end_day": "현재",
  "description": "회사명, 부서, 담당업무, 보유기술, 성과를 포함한 상세 설명"
}

규칙:
- 최근 경력부터 나열 (역순)
- 일자를 모르면 "01"로 채우세요
- "현재" 근무 중이면 end_year/month/day에 "현재"를 넣으세요
- description은 채용공고에 맞게 작성된 이력서 내용을 그대로 사용하세요
- 반드시 JSON 배열만 응답하세요"""


async def fill_table_template(
    template_path: str | Path,
    resume: TailoredResume,
    output_path: str | Path,
    llm: LLMClient | None = None,
) -> Path:
    """Fill a table-based DOCX template with resume experience data.

    If an LLM client is provided, it extracts structured experience data
    from the resume. Otherwise, it fills the description column with
    the experience section content directly.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("DOCX 파일에 표가 없습니다.")

    table = doc.tables[0]

    # Find the first data row (skip header rows)
    data_start_row = _find_data_start_row(table)

    if llm:
        entries = await _extract_entries_with_llm(llm, resume)
    else:
        entries = _extract_entries_simple(resume)

    # Detect column layout from the header row
    col_layout = _detect_column_layout(table, data_start_row)

    # Fill table rows
    available_rows = len(table.rows) - data_start_row
    for i, entry in enumerate(entries):
        if i >= available_rows:
            _add_table_row(table)

        row_idx = data_start_row + i
        row = table.rows[row_idx]
        _fill_row(row, entry, col_layout)

    doc.save(str(output_path))
    return output_path


def fill_table_template_sync(
    template_path: str | Path,
    resume: TailoredResume,
    output_path: str | Path,
    llm: LLMClient | None = None,
) -> Path:
    """Synchronous wrapper for fill_table_template."""
    return asyncio.run(
        fill_table_template(template_path, resume, output_path, llm)
    )


def _find_data_start_row(table) -> int:
    """Find the first empty row (data row) in the table."""
    for i, row in enumerate(table.rows):
        all_empty = all(not cell.text.strip() for cell in row.cells)
        if all_empty:
            return i
    return min(2, len(table.rows))


def _detect_column_layout(table, data_start_row: int) -> str:
    """Detect the column layout pattern from header rows.

    Returns:
        "ymd_sep" — 년|월|일|~|년|월|일|description  (8 cols, date cells separate)
        "period_desc" — period_merged|description      (2+ unique cols)
        "single" — everything in one cell
    """
    # Check the sub-header row (typically row 1) for 년/월/일 pattern
    if data_start_row >= 2:
        header_row = table.rows[data_start_row - 1]
        texts = [cell.text.strip() for cell in header_row.cells]
        # Look for the 년|월|일|~|년|월|일 pattern
        if len(texts) >= 7 and texts[0] == "년" and texts[3] == "~":
            return "ymd_sep"

    # Check data row cell count
    if data_start_row < len(table.rows):
        unique = _get_unique_cells(table.rows[data_start_row])
        if len(unique) == 1:
            return "single"

    return "period_desc"


def _fill_row(row, entry: dict, col_layout: str) -> None:
    """Fill a single table row based on the detected column layout."""
    cells = row.cells

    if col_layout == "ymd_sep" and len(cells) >= 8:
        # 년 | 월 | 일 | ~ | 년 | 월 | 일 | description
        _set_cell_text(cells[0], entry.get("start_year", ""))
        _set_cell_text(cells[1], entry.get("start_month", ""))
        _set_cell_text(cells[2], entry.get("start_day", ""))
        _set_cell_text(cells[3], "~")
        end_year = entry.get("end_year", "")
        if end_year == "현재":
            _set_cell_text(cells[4], "현재")
            _set_cell_text(cells[5], "")
            _set_cell_text(cells[6], "")
        else:
            _set_cell_text(cells[4], end_year)
            _set_cell_text(cells[5], entry.get("end_month", ""))
            _set_cell_text(cells[6], entry.get("end_day", ""))
        _set_cell_text(cells[7], entry.get("description", ""))

    elif col_layout == "period_desc":
        unique = _get_unique_cells(row)
        if len(unique) >= 2:
            period_text = _format_period(entry)
            _set_cell_text(unique[0], period_text)
            _set_cell_text(unique[-1], entry.get("description", ""))
        elif len(unique) == 1:
            text = f"{_format_period(entry)}\n\n{entry.get('description', '')}"
            _set_cell_text(unique[0], text)

    else:  # single
        unique = _get_unique_cells(row)
        text = f"{_format_period(entry)}\n\n{entry.get('description', '')}"
        _set_cell_text(unique[0], text)


def _get_unique_cells(row) -> list:
    """Get unique cells in a row (skip duplicates from merged cells)."""
    seen = set()
    unique = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id not in seen:
            seen.add(cell_id)
            unique.append(cell)
    return unique


def _set_cell_text(cell, text: str, font_size: int = 10) -> None:
    """Set cell text, preserving existing formatting where possible."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""

    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = text
        else:
            run = para.add_run(text)
            run.font.size = Pt(font_size)
    else:
        cell.add_paragraph(text)


def _format_period(entry: dict) -> str:
    """Format a period string from entry data."""
    start = f"{entry.get('start_year', '')}.{entry.get('start_month', '')}.{entry.get('start_day', '')}"
    end_year = entry.get("end_year", "")
    if end_year == "현재":
        end = "현재"
    else:
        end = f"{end_year}.{entry.get('end_month', '')}.{entry.get('end_day', '')}"
    return f"{start} ~ {end}"


def _add_table_row(table) -> None:
    """Add a new row to the table, copying format from the last row."""
    from copy import deepcopy
    from lxml import etree

    last_row = table.rows[-1]
    new_tr = deepcopy(last_row._tr)
    # Clear text content in new row
    for tc in new_tr.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
        tc.text = ""
    table._tbl.append(new_tr)


async def _extract_entries_with_llm(llm: LLMClient, resume: TailoredResume) -> list[dict]:
    """Use LLM to extract structured experience entries from resume."""
    prompt = f"""다음 이력서에서 경력/프로젝트 항목을 추출하세요.
최근 경력부터 나열하세요.

이력서:
{resume.full_markdown}

JSON 배열로만 응답하세요."""

    data = await llm.generate_json(
        prompt=prompt,
        system=EXTRACT_PROMPT_SYSTEM,
        model="claude-haiku-4-5-20251001",
    )

    # Handle case where LLM returns {"entries": [...]} instead of [...]
    if isinstance(data, dict):
        for key in ("entries", "items", "experiences", "경력"):
            if key in data:
                return data[key]
        return list(data.values())[0] if data else []

    return data if isinstance(data, list) else []


def _extract_entries_simple(resume: TailoredResume) -> list[dict]:
    """Extract experience entries without LLM (basic parsing)."""
    experience_section = None
    for section in resume.sections:
        if section.id in ("experience", "경력사항") or "경력" in section.label:
            experience_section = section
            break

    if not experience_section:
        return [{"description": resume.full_markdown}]

    # Simple split by ### or --- markers
    content = experience_section.content
    entries = []

    # Split by ### headings (company names)
    parts = re.split(r"(?=^### )", content, flags=re.MULTILINE)
    for part in parts:
        part = part.strip()
        if not part:
            continue

        # Try to extract dates
        date_match = re.search(r"(\d{4})\.(\d{2})\s*~\s*(현재|\d{4})\.?(\d{2})?", part)
        entry = {
            "start_year": date_match.group(1) if date_match else "",
            "start_month": date_match.group(2) if date_match else "",
            "start_day": "01",
            "end_year": date_match.group(3) if date_match else "",
            "end_month": date_match.group(4) if date_match else "",
            "end_day": "현재" if date_match and date_match.group(3) == "현재" else "01",
            "description": _strip_md_for_docx(part),
        }
        entries.append(entry)

    return entries or [{"description": _strip_md_for_docx(content)}]


def _strip_md_for_docx(text: str) -> str:
    """Strip markdown formatting for DOCX content."""
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
